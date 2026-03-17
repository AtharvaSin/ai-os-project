#!/usr/bin/env python3
"""
generate_docx_templates.py
Creates three branded .docx templates — one per brand context.
Outputs to /tmp/brand-templates/ for upload to Drive.

Usage:
  pip install python-docx --break-system-packages
  python3 generate_docx_templates.py

Output files:
  /tmp/brand-templates/ai_os_template.docx       (Context A)
  /tmp/brand-templates/bharatvarsh_template.docx  (Context B)
  /tmp/brand-templates/portfolio_template.docx    (Context C)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = '/tmp/brand-templates'


# ──────────────────────────────────────────────────────
# Shared utilities
# ──────────────────────────────────────────────────────

def hex_to_rgb(hex_color):
    """Convert hex string to (r, g, b) tuple."""
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_font(run, name, size_pt, bold=False, italic=False, color_hex=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)


def set_paragraph_spacing(para, before_pt=0, after_pt=6, line_spacing_pt=None):
    para.paragraph_format.space_before = Pt(before_pt)
    para.paragraph_format.space_after = Pt(after_pt)
    if line_spacing_pt:
        para.paragraph_format.line_spacing = Pt(line_spacing_pt)


def add_colored_heading(doc, text, level, font_name, font_size,
                         color_hex, bold=True, after_pt=6, before_pt=12):
    """Add a heading paragraph with specific font/color."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before_pt=before_pt, after_pt=after_pt)
    run = para.add_run(text)
    set_font(run, font_name, font_size, bold=bold, color_hex=color_hex)
    return para


def add_body_text(doc, text, font_name, font_size, color_hex, after_pt=4):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, after_pt=after_pt)
    run = para.add_run(text)
    set_font(run, font_name, font_size, color_hex=color_hex)
    return para


def add_horizontal_rule(doc, color_hex):
    """Add a styled horizontal rule paragraph."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before_pt=4, after_pt=4)
    run = para.add_run('─' * 72)
    set_font(run, 'Courier New', 8, color_hex=color_hex)
    return para


def add_label(doc, text, font_name, font_size, color_hex):
    """Add an uppercase label (section eyebrow)."""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before_pt=16, after_pt=2)
    run = para.add_run(text.upper())
    set_font(run, font_name, font_size, bold=True, color_hex=color_hex)
    from docx.oxml import OxmlElement
    # Letter spacing via XML (approximate)
    return para


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    """Set page margins in inches."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


# ──────────────────────────────────────────────────────
# Context A: AI OS System
# ──────────────────────────────────────────────────────

def create_context_a_template(output_path):
    """
    AI OS System template.
    Dark professional. DM Sans + JetBrains Mono. Electric teal accent.
    Note: Word has limited dark theme support — this template uses
    the closest achievable approximation with white page + dark accents.
    The full dark experience requires HTML/PDF rendering.
    """
    print("Creating Context A (AI OS System) template...")
    doc = Document()
    set_page_margins(doc)
    
    # Accent color — update from BRAND_IDENTITY.md after extraction
    ACCENT = '#00E5CC'      # Electric teal — placeholder until extraction
    TEXT_PRIMARY = '#0d0d14'  # Dark text (since Word is light-bg)
    TEXT_SECONDARY = '#404040'
    TEXT_MUTED = '#737373'
    HEADING_COLOR = '#0d0d14'
    
    # ── Cover page content ──
    doc.add_paragraph()  # spacer
    
    label_para = doc.add_paragraph()
    label_run = label_para.add_run('AI OPERATING SYSTEM — DOCUMENT')
    set_font(label_run, 'DM Sans', 9, bold=True, color_hex=ACCENT)
    set_paragraph_spacing(label_para, before_pt=72, after_pt=4)
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Document Title')
    set_font(title_run, 'Instrument Serif', 32, color_hex=HEADING_COLOR)
    set_paragraph_spacing(title_para, before_pt=8, after_pt=16)
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run('Subtitle or classification line')
    set_font(subtitle_run, 'DM Sans', 14, color_hex=TEXT_SECONDARY)
    set_paragraph_spacing(subtitle_para, after_pt=40)
    
    add_horizontal_rule(doc, ACCENT)
    
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run('Author: Atharva Singh     Date: [DATE]     Version: 1.0')
    set_font(meta_run, 'JetBrains Mono', 9, color_hex=TEXT_MUTED)
    set_paragraph_spacing(meta_para, before_pt=8, after_pt=4)
    
    doc.add_page_break()
    
    # ── Section structure demo ──
    add_label(doc, '01 — Overview', 'DM Sans', 10, ACCENT)
    add_colored_heading(doc, 'Section Heading (DM Sans 700)', 1, 'DM Sans', 18, HEADING_COLOR)
    add_body_text(doc, 
        'Body text uses DM Sans 400 at 11pt. This template establishes the Context A '
        'design language for the AI Operating System. Use this template for: PRDs, '
        'decision memos, research reports, system documentation, and sprint reviews.',
        'DM Sans', 11, TEXT_SECONDARY, after_pt=8)
    
    add_label(doc, '02 — Usage Guide', 'DM Sans', 10, ACCENT)
    add_colored_heading(doc, 'How to Use This Template', 2, 'DM Sans', 14, HEADING_COLOR)
    add_body_text(doc,
        'Replace all placeholder text. The numbered section labels (01 —, 02 —, etc.) '
        'are the structural grammar of Context A. Keep them. Use JetBrains Mono for all '
        'data values, metrics, version numbers, and technical identifiers.',
        'DM Sans', 11, TEXT_SECONDARY)
    
    # ── Data label demo ──
    data_para = doc.add_paragraph()
    data_run = data_para.add_run('METRIC: 34 tools    STATUS: operational    SPRINT: 8')
    set_font(data_run, 'JetBrains Mono', 10, color_hex=TEXT_MUTED)
    set_paragraph_spacing(data_para, before_pt=8, after_pt=16)
    
    add_horizontal_rule(doc, '#1f1f35')
    
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run(
        'Context A — AI OS System  |  Brand: Obsidian Aurora / Wibify-inspired  |  '
        'Fonts: Instrument Serif, DM Sans, JetBrains Mono  |  Accent: Electric Teal'
    )
    set_font(caption_run, 'JetBrains Mono', 8, color_hex=TEXT_MUTED)
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(output_path)
    print(f"  ✅ Saved: {output_path}")


# ──────────────────────────────────────────────────────
# Context B: Bharatvarsh
# ──────────────────────────────────────────────────────

def create_context_b_template(output_path):
    """
    Bharatvarsh template.
    Cinematic-military. Bebas Neue + Crimson Pro + Inter.
    Mustard gold accent. Navy/obsidian atmosphere.
    """
    print("Creating Context B (Bharatvarsh) template...")
    doc = Document()
    set_page_margins(doc)
    
    MUSTARD = '#B8921F'     # Darker mustard for print readability
    NAVY = '#0B2742'
    TEXT_MAIN = '#1A1A2A'
    TEXT_SECONDARY = '#4A4A5A'
    TEXT_MUTED = '#6A6A7A'
    
    # ── Cover page ──
    doc.add_paragraph()
    
    top_label = doc.add_paragraph()
    top_run = top_label.add_run('CLASSIFIED — BHARATVARSH INTELLIGENCE FILE')
    set_font(top_run, 'JetBrains Mono', 8, bold=True, color_hex=MUSTARD)
    set_paragraph_spacing(top_label, before_pt=72, after_pt=24)
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('DOCUMENT TITLE')
    # Bebas Neue approximation — fallback to Impact or Arial Narrow
    set_font(title_run, 'Bebas Neue', 42, color_hex=NAVY)
    set_paragraph_spacing(title_para, before_pt=4, after_pt=8)
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run('Subtitle — In Crimson Pro for narrative weight')
    set_font(subtitle_run, 'Crimson Pro', 16, italic=True, color_hex=TEXT_SECONDARY)
    set_paragraph_spacing(subtitle_para, after_pt=32)
    
    rule_para = doc.add_paragraph()
    rule_run = rule_para.add_run('━' * 60)
    set_font(rule_run, 'JetBrains Mono', 9, color_hex=MUSTARD)
    set_paragraph_spacing(rule_para, before_pt=4, after_pt=8)
    
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run('AUTHOR: ATHARVA SINGH     DATE: [DATE]     CLASS: CONFIDENTIAL')
    set_font(meta_run, 'JetBrains Mono', 8, color_hex=TEXT_MUTED)
    
    doc.add_page_break()
    
    # ── Content demo ──
    faction_label = doc.add_paragraph()
    faction_run = faction_label.add_run('// FACTION: REPUBLIC OF BHARAT')
    set_font(faction_run, 'JetBrains Mono', 9, bold=True, color_hex=MUSTARD)
    set_paragraph_spacing(faction_label, before_pt=16, after_pt=4)
    
    section_head = doc.add_paragraph()
    head_run = section_head.add_run('SECTION HEADING IN BEBAS NEUE')
    set_font(head_run, 'Bebas Neue', 22, color_hex=NAVY)
    set_paragraph_spacing(section_head, before_pt=4, after_pt=6)
    
    body_para = doc.add_paragraph()
    body_run = body_para.add_run(
        'Narrative body text uses Crimson Pro at 12pt for a literary, historically-weighted '
        'feel. This is appropriate for lore documents, character profiles, timeline entries, '
        'and any content that needs the Bharatvarsh world to feel real and lived-in.'
    )
    set_font(body_run, 'Crimson Pro', 12, color_hex=TEXT_MAIN)
    set_paragraph_spacing(body_para, after_pt=10)
    
    ui_para = doc.add_paragraph()
    ui_run = ui_para.add_run(
        'Standard UI copy, briefings, and operational text use Inter at 11pt. '
        'Clear, readable, military-functional.'
    )
    set_font(ui_run, 'Inter', 11, color_hex=TEXT_SECONDARY)
    set_paragraph_spacing(ui_para, after_pt=16)
    
    add_horizontal_rule(doc, MUSTARD)
    
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run(
        'Context B — Bharatvarsh  |  Fonts: Bebas Neue, Crimson Pro, Inter, JetBrains Mono  '
        '|  Accent: Mustard Gold #F1C232  |  Navy #0B2742'
    )
    set_font(caption_run, 'JetBrains Mono', 8, color_hex=TEXT_MUTED)
    
    doc.save(output_path)
    print(f"  ✅ Saved: {output_path}")


# ──────────────────────────────────────────────────────
# Context C: Portfolio
# ──────────────────────────────────────────────────────

def create_context_c_template(output_path):
    """
    Portfolio template.
    Clean, dual-identity. Inter universal. Violet primary, coral accent.
    """
    print("Creating Context C (Portfolio) template...")
    doc = Document()
    set_page_margins(doc)
    
    VIOLET = '#7c3aed'
    CORAL = '#ea580c'
    TEXT_HEADING = '#171717'
    TEXT_BODY = '#404040'
    TEXT_MUTED = '#737373'
    
    # ── Cover page ──
    doc.add_paragraph()
    
    label_para = doc.add_paragraph()
    label_run = label_para.add_run('ATHARVA SINGH — PORTFOLIO')
    set_font(label_run, 'Inter', 9, bold=True, color_hex=VIOLET)
    set_paragraph_spacing(label_para, before_pt=72, after_pt=8)
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Document Title')
    set_font(title_run, 'Inter', 36, bold=True, color_hex=TEXT_HEADING)
    set_paragraph_spacing(title_para, before_pt=4, after_pt=8)
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run('Subtitle — clean, readable, professional')
    set_font(subtitle_run, 'Inter', 14, color_hex=TEXT_BODY)
    set_paragraph_spacing(subtitle_para, after_pt=32)
    
    rule_para = doc.add_paragraph()
    rule_run = rule_para.add_run('─' * 64)
    set_font(rule_run, 'Inter', 9, color_hex='#e5e5e5')
    set_paragraph_spacing(rule_para, before_pt=4, after_pt=8)
    
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run('Atharva Singh  ·  atharvasingh.com  ·  [DATE]')
    set_font(meta_run, 'Inter', 10, color_hex=TEXT_MUTED)
    
    doc.add_page_break()
    
    # ── Content demo ──
    add_colored_heading(doc, 'Section Heading (Inter 700, 18pt)', 1,
                        'Inter', 18, TEXT_HEADING, bold=True)
    
    body_para = doc.add_paragraph()
    body_run = body_para.add_run(
        'Portfolio template body uses Inter at 11pt throughout — light mode, clean, '
        'highly readable. This template suits professional case studies, CV documents, '
        'cover letters, personal branding materials, and LinkedIn-adjacent content.'
    )
    set_font(body_run, 'Inter', 11, color_hex=TEXT_BODY)
    set_paragraph_spacing(body_para, after_pt=8)
    
    add_colored_heading(doc, 'Subsection (Inter 600, 13pt, Violet)', 2,
                        'Inter', 13, VIOLET, bold=True)
    
    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(
        'Use violet (#8b5cf6 → #7c3aed) for primary structure and navigation. '
        'Use coral/orange (#f97316 → #ea580c) for callouts and creative emphasis. '
        'Never mix these with Mustard or Electric Teal from other contexts.'
    )
    set_font(sub_run, 'Inter', 11, color_hex=TEXT_BODY)
    set_paragraph_spacing(sub_para, after_pt=8)
    
    callout_para = doc.add_paragraph()
    callout_run = callout_para.add_run('→ Creative callout in coral: use sparingly for emphasis')
    set_font(callout_run, 'Inter', 11, bold=True, color_hex=CORAL)
    set_paragraph_spacing(callout_para, before_pt=8, after_pt=16)
    
    add_horizontal_rule(doc, '#e5e5e5')
    
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run(
        'Context C — Portfolio  |  Font: Inter (all weights)  |  '
        'Primary: Violet #8b5cf6  |  Accent: Coral #f97316'
    )
    set_font(caption_run, 'Inter', 8, color_hex=TEXT_MUTED)
    
    doc.save(output_path)
    print(f"  ✅ Saved: {output_path}")


# ──────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────

def main():
    print("Generating brand document templates...\n")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    create_context_a_template(os.path.join(OUTPUT_DIR, 'ai_os_template.docx'))
    create_context_b_template(os.path.join(OUTPUT_DIR, 'bharatvarsh_template.docx'))
    create_context_c_template(os.path.join(OUTPUT_DIR, 'portfolio_template.docx'))
    
    print(f"\n✅ All templates generated in: {OUTPUT_DIR}")
    print("\nNext steps:")
    print("1. Review templates by downloading and opening in Word/Docs")
    print("2. If Bebas Neue / DM Sans / Instrument Serif aren't installed locally,")
    print("   they'll fall back to system fonts in Word preview — this is expected.")
    print("   The font NAMES are embedded and will display correctly where fonts are installed.")
    print("3. Upload to Drive: AI OS/BRAND_TEMPLATES/ using AIOSMCP upload_file tool")
    print("   - ai_os_template.docx → context-a-ai-os/")
    print("   - bharatvarsh_template.docx → context-b-bharatvarsh/")
    print("   - portfolio_template.docx → context-c-portfolio/")


if __name__ == '__main__':
    main()
